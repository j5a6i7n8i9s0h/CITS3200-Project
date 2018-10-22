from docx.api import Document
import json, sys

def run(coversheet_data, name):
    document = Document('py-docx/test.docx')
    i=0
    j=0
    dictionary = coversheet_data['contact_details']
    for rows in document.tables[0].rows:
        head = rows.cells[0].text
        if head in dictionary:
            if head in ['Protocol Title :', 'Monitoring Start Date :']:
                rows.cells[1].paragraphs[0].add_run(dictionary[head]).bold = True
            else: 
                if 'Monitor 3' in head:
                    if i==0:
                        rows.cells[1].paragraphs[0].add_run(dictionary[head][0]).bold = True
                        rows.cells[2].paragraphs[0].add_run(dictionary[head][1]).bold = True
                    else: 
                        rows.cells[1].paragraphs[0].add_run(" - Ticked " if  dictionary['Supervisor :'] else " - Unticked").bold = True
                    i+=1
                else:
                    rows.cells[1].paragraphs[0].add_run(dictionary[head][0]).bold = True
                    rows.cells[2].paragraphs[0].add_run(dictionary[head][1]).bold = True
    document.tables[1].rows[0].cells[0].paragraphs[0].add_run(coversheet_data['species_phenotype_issues']['Species']).bold = True
    document.save('animalwellbeing/static/animalwellbeing/coversheets/{}.docx'.format(name))

