import docx
import datetime
from datetime import timedelta
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json, sys

coversheet_data = json.loads(sys.argv[1])
contact_details = coversheet_data['contact_details']
species_phenotype_issues = coversheet_data['species_phenotype_issues']
monitoring_criteria = coversheet_data['monitoring_criteria']
monitoring_frequency = coversheet_data['monitoring_frequency']
aec = coversheet_data['aec']

recording_sheet = docx.Document("py-docx/rec_sheet.docx")

table_one = recording_sheet.tables[0]
table_two = recording_sheet.tables[1]
table_three = recording_sheet.tables[2]
table_four = recording_sheet.tables[3]
table_five = recording_sheet.tables[4]

def insert_start_date(year, month, day):
    global rec_start_date
    rec_start_date = datetime.datetime(year, month, day)
    str_date = rec_start_date.strftime('%d/%m/%y')
    recording_sheet.paragraphs[0].runs[15].add_text(str_date)
    return

def insert_weighing_freq(weighing_freq):
    table_one.rows[1].cells[5].paragraphs[0].add_run(weighing_freq)
    return

def insert_monitoring_freq(monitoring_freq):
    table_one.rows[1].cells[7].paragraphs[0].add_run(monitoring_freq)
    return

def insert_weight_dates(freq):
    start_date = datetime.datetime.now()
    alpha = start_date - rec_start_date
    while (alpha.days % freq != 0):
        start_date = start_date + timedelta(days = 1)
        alpha = start_date - rec_start_date
    str_date = start_date.strftime('%d/%m/%y')
    a = 2
    for x in range(7):
        table_two.rows[0].cells[a].paragraphs[0].add_run(start_date.strftime('%d/%m/%y'))
        start_date = start_date + timedelta(days = freq)
        a += 2
    return

def insert_monitoring_dates(freq):
    start_date = datetime.datetime.now()
    alpha = start_date - rec_start_date
    while (alpha.days % freq != 0):
        start_date = start_date + timedelta(days = 1)
        alpha = start_date - rec_start_date
    str_date = start_date.strftime('%d/%m/%y')
    for x in range(7):
        delta = start_date - rec_start_date
        day_no = (delta.days/freq) + 1
        table_three.rows[1].cells[x+2].paragraphs[0].add_run('Day ' + str(int(day_no))).bold = True
        table_three.rows[2].cells[x+2].paragraphs[0].add_run(start_date.strftime('%d/%m/%y'))
        start_date = start_date + timedelta(days = freq)
    return

def insert_aec_protocol(protocol):
    table_one.rows[1].cells[1].paragraphs[0].runs[0].add_text(protocol)

def insert_weight_threshold(w_th1, w_th2):
    table_one.rows[3].cells[5].paragraphs[0].runs[0].add_text(w_th1 + '%')
    table_one.rows[3].cells[6].paragraphs[0].runs[0].add_text(w_th2 + '%')

def insert_tumor_threshold(t_th1, t_th2):
    table_one.rows[3].cells[7].paragraphs[0].runs[0].add_text(t_th1)
    table_one.rows[3].cells[8].paragraphs[0].runs[0].add_text(t_th2)

def insert_species(str_species):
    table_one.rows[1].cells[3].paragraphs[0].add_run(str_species)
    return

def insert_criteria(std_criteria, proj_criteria):
    list_criteria = std_criteria.split('#')
    row = 1
    for x in range(len(list_criteria)):
        templist = list_criteria[x].split('@')
        table_four.add_row()
        table_four.rows[row - 1].cells[0].merge(table_four.rows[(x+1)].cells[0])
        table_four.rows[row].cells[1].paragraphs[0].add_run(templist[0])
        row += 1
    list_criteria = proj_criteria.split('#')
    for x in range(len(list_criteria)):
        templist = list_criteria[x].split('@')
        table_four.add_row()
        table_four.rows[row - 1].cells[0].merge(table_four.rows[(row)].cells[0])
        table_four.rows[row].cells[1].paragraphs[0].add_run(templist[0])
        row += 1


if 'Monitoring Start Date :' in contact_details:
    start_date = contact_details['Monitoring Start Date :'].split('-')
    if (start_date[0].isdigit() == start_date[1].isdigit() == start_date[2].isdigit() == True):
        insert_start_date(int(start_date[0]), int(start_date[1]), int(start_date[2]))

        if 'weighing_frequency' in monitoring_frequency:
            insert_weighing_freq(monitoring_frequency['weighing_frequency'])
            weigh_freq = monitoring_frequency['weighing_frequency'].split(' ')
            if (weigh_freq[2].isdigit()):
                insert_weight_dates(int(weigh_freq[2]))

        if 'monitoring_frequency' in monitoring_frequency:
            insert_monitoring_freq(monitoring_frequency['monitoring_frequency'])
            monitor_freq = monitoring_frequency['monitoring_frequency'].split(' ')
            if (monitor_freq[2].isdigit()):
                insert_monitoring_dates(int(monitor_freq[2]))

else:
    if 'weighing_frequency' in monitoring_frequency:
        insert_weighing_freq(monitoring_frequency['weighing_frequency'])

    if 'monitoring_frequency' in monitoring_frequency:
        insert_monitoring_freq(monitoring_frequency['monitoring_frequency'])

#if 'Protocol Title :' in contact_details:
#    insert_aec_protocol(contact_details['Protocol Title :'])

table_one.rows[3].cells[5].paragraphs[0].runs[0].add_text(aec['aec1'] + '%')
table_one.rows[3].cells[6].paragraphs[0].runs[0].add_text(aec['aec4'] + '%')
table_one.rows[3].cells[7].paragraphs[0].runs[0].add_text(aec['aec5'])
table_one.rows[3].cells[8].paragraphs[0].runs[0].add_text(aec['aec8'])

if 'Species' in species_phenotype_issues:
    insert_species(species_phenotype_issues['Species'])

if 'standard_criteria' in monitoring_criteria:
    if 'project_criteria' in monitoring_criteria:
        insert_criteria(monitoring_criteria['standard_criteria'], monitoring_criteria['project_criteria'])

recording_sheet.save('animalwellbeing/static/animalwellbeing/recordingsheets/{}.docx'.format(sys.argv[2]))
