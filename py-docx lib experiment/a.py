from docx import Document 


document = Document('test.docx')

document.tables[1].rows[0].cells[0].paragraphs[0].add_run('BRUUUUHHHH').bold = True

#document.tables[3].rows[0].cells[0].paragraphs[0].add_run('dumbledore dies in hewlett packet').bold = True

''' 
 0 table 1: contact details 
 1 table 2: SPECIES / PHENOTYPE / MODEL ISSUES
 2 table 3: MONITORING CRITERIA AND SCORING 
 3 table 4: MONITORING FREQUENCY
 4 table 5: Montiroing Freq : Anathesia , general etc
 5 table 6:	5) ACTIONS AND INTERVENTIONS
 6 table 7: 6)	AEC INTERVENTIONS for Body Weight Loss and Subcutaneous Tumour Size     (as appropriate to the project) 
 7 table 8: 7)	INSTRUCTIONS for the conduct of the monitoring and recording

'''
#print (type(document.tables[0].rows.height))
# for count,val in enumerate(document.tables[0].rows):
#     print count,val
dictionary ={
    'Protocol Title :' : 'ironman',
    'Monitoring Start Date :' : 'today' , 
    'Chief Investigator :' : ['kanye_west' , '123123'],
    'Emergency Contact :' : ['thor', ' 123123141342'],
    'Monitor 1 :' : ['captainamerica', ' 123123141342'],
    'Monitor 2 :' : ['lilpump', ' 123123141342'],
    'Monitor 3 :' : ['thanos ', ' 12342'],
    'Person responsible for euthanasia :' : ['h3h3', ' 12'],
    'Other experts :' : ['voldemort', ' 2'],
    
}

standard_dict ={
    'acitivty': '1',
    'jumpy jacks' : '2', 
    'racks on racks on  racks' : '0', 
    'justin bieber': '1',
}

specifc_dict={
    'asd' :'2', 
    'avf': '1', 
    'sa': '0',
}
i=0
for rows in document.tables[0].rows:
        head = rows.cells[0].text
        if head in dictionary:
            if head in ['Protocol Title :', 'Monitoring Start Date :']:
                rows.cells[1].paragraphs[0].add_run(dictionary[head]).bold = True
            else: 
                rows.cells[2].paragraphs[0].add_run(dictionary[head][1]).bold = True
                if 'Monitor 3' in head:
                    if i==0:
                        rows.cells[1].paragraphs[0].add_run(dictionary[head][0]).bold = True
                    i+=1
                else:
                    rows.cells[1].paragraphs[0].add_run(dictionary[head][0]).bold = True
criteras = len(standard_dict)
proj_specific_critera = len(specifc_dict)
total = 0
if criteras > 4: 
    for i in range(criteras-4): 
        document.tables[2].add_row()
        document.tables[2].style = 'Table Grid'
j=1
for key in standard_dict: 
    document.tables[2].rows[j].cells[0].paragraphs[0].add_run(key).bold = True   #.cells[0].key 
    document.tables[2].rows[j].cells[int(standard_dict[key])+1].paragraphs[0].add_run('this').bold = True   #.cells[0].key 
    total += int(standard_dict[key])
    j+=1

if proj_specific_critera > 2: 
        for i in range(proj_specific_critera-2): 
            document.tables[3].add_row()
            document.tables[3].style = 'Table Grid'
j=1
for key in specifc_dict: 
    document.tables[3].rows[j].cells[0].paragraphs[0].add_run(key).bold = True   #.cells[0].key 
    document.tables[3].rows[j].cells[int(specifc_dict[key])+1].paragraphs[0].add_run('this').bold = True   #.cells[0].key 
    total += int(specifc_dict[key])
    j+=1
#modifyBorder(document.tables[2])
document.tables[5].rows[0].cells[0].paragraphs[0].add_run('good_car').bold = True


type_of_sheet = '3'
''' processing  type of recording sheet '''
string = document.tables[6].rows[1].cells[int(type_of_sheet)].text
string = string[:string.index("[")] + '[x]'

document.tables[6].rows[1].cells[int(type_of_sheet)].paragraphs[0].text = string
if type_of_sheet=='3' : 
    document.tables[6].rows[1].cells[int(type_of_sheet)].paragraphs[1].text = 'Specify : {}'.format('test_')


from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

shading_elm_1 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
shading_elm_2 = parse_xml(r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
#table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)


if total==0:
    document.tables[7].rows[1].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    document.tables[7].rows[1].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)
elif total==1: 
    document.tables[7].rows[2].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    document.tables[7].rows[2].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2) 
elif total>=2 and total<=4: 
    document.tables[7].rows[3].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    document.tables[7].rows[3].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)
else:
    document.tables[7].rows[4].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    document.tables[7].rows[4].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)

document.save('test2.docx')
